import io
import os
import time
from typing import Optional, Tuple
from utils_ui import render_title

import requests
import streamlit as st

# ==========================================================
# PÁGINA: Transcripción (Audio → Texto)
# - Compatible con el proyecto base (login en app.py)
# - SIN st.set_page_config() (solo en app.py)
# - SIN st.switch_page() / st.rerun() (evita loops)
# - Exporta TXT siempre; DOCX/PDF si tienes dependencias
# ==========================================================

# --------- Login guard (misma llave que app.py del proyecto base)
def require_login() -> None:
    if not st.session_state.get("auth", False):
        st.error("🔒 Inicia sesión para usar esta herramienta.")
        if hasattr(st, "page_link"):
            st.page_link("app.py", label="Ir al Login", icon="🔐", use_container_width=True)
        st.stop()


require_login()

# --------- UI Header
st.markdown("# 🎧 Transcripción")
st.caption("Carga un audio y obtén la transcripción en español (lista para copiar o exportar).")

if hasattr(st, "page_link"):
    st.page_link("app.py", label="⬅️ Volver al Portafolio", icon="🏠", use_container_width=True)

st.divider()

with st.expander("🔒 Privacidad (cómo funciona)", expanded=False):
    st.write(
        "- El audio se envía al servicio de transcripción y se devuelve el texto.\n"
        "- Esta app no guarda el archivo ni la transcripción.\n"
        "- En Streamlit Cloud, el archivo vive solo en memoria durante la ejecución."
    )

# --------- Config
MODEL = "gpt-4o-mini-transcribe"
LANGUAGE = "es"
PROMPT = (
    "Transcribe únicamente en español. Contexto industrial/operativo. "
    "Conserva siglas y términos técnicos. Mantén números, unidades y nombres propios.\n"
    "Términos frecuentes: IMEMSA, gel coat, fibra de vidrio, infusión al vacío, T-top, quilla, patín, borda, consola."
)

MAX_FILE_MB = 25


# ==========================================================
# OpenAI Transcription via HTTP (estable, sin depender del SDK)
# Docs: POST /audio/transcriptions
# ==========================================================
def _get_openai_api_key() -> Optional[str]:
    # Prioriza secrets (Streamlit Cloud) y luego variables de entorno
    return (
        st.secrets.get("OPENAI_API_KEY") if hasattr(st, "secrets") else None
    ) or os.getenv("OPENAI_API_KEY")


def _guess_mime(filename: str) -> str:
    name = filename.lower()
    if name.endswith(".mp3"):
        return "audio/mpeg"
    if name.endswith(".m4a"):
        return "audio/mp4"
    if name.endswith(".wav"):
        return "audio/wav"
    if name.endswith(".webm"):
        return "audio/webm"
    if name.endswith(".mp4"):
        return "video/mp4"
    if name.endswith(".mpeg") or name.endswith(".mpg"):
        return "video/mpeg"
    if name.endswith(".mpga"):
        return "audio/mpeg"
    return "application/octet-stream"


def transcribe_openai(audio_bytes: bytes, filename: str) -> Tuple[str, Optional[dict]]:
    api_key = _get_openai_api_key()
    if not api_key:
        st.error("Falta la API Key.")
        st.info(
            "En Streamlit Cloud agrega un Secret llamado **OPENAI_API_KEY** "
            "o define la variable de entorno `OPENAI_API_KEY`."
        )
        st.stop()

    url = "https://api.openai.com/v1/audio/transcriptions"
    headers = {"Authorization": f"Bearer {api_key}"}

    files = {
        "file": (filename, io.BytesIO(audio_bytes), _guess_mime(filename)),
    }
    data = {
        "model": MODEL,
        "language": LANGUAGE,
        "prompt": PROMPT,
        "response_format": "json",  # gpt-4o-mini-transcribe soporta json
    }

    resp = requests.post(url, headers=headers, files=files, data=data, timeout=180)

    if resp.status_code >= 400:
        # Muestra error con contexto (sin filtrar key)
        st.error(f"Error en transcripción (HTTP {resp.status_code}).")
        try:
            st.code(resp.json())
        except Exception:
            st.code(resp.text[:2000])
        st.stop()

    payload = resp.json()
    text = payload.get("text", "").strip()
    if not text:
        st.warning("La respuesta no trajo texto. Revisa si el audio tiene voz clara.")
    return text, payload


# ==========================================================
# Export helpers (opcionales)
# ==========================================================
def to_docx_bytes(title: str, body: str) -> Optional[bytes]:
    try:
        from docx import Document
    except Exception:
        return None

    doc = Document()
    doc.add_heading(title, level=1)
    for para in body.split("\n"):
        doc.add_paragraph(para)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def to_pdf_bytes(title: str, body: str) -> Optional[bytes]:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas
    except Exception:
        return None

    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=LETTER)
    width, height = LETTER

    # Simple layout
    x = 50
    y = height - 60
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, title)
    y -= 25
    c.setFont("Helvetica", 11)

    # Wrap lines
    max_chars = 95
    lines = []
    for raw in body.split("\n"):
        raw = raw.rstrip()
        if not raw:
            lines.append("")
            continue
        while len(raw) > max_chars:
            cut = raw[:max_chars]
            # intenta cortar por espacio
            if " " in cut:
                cut = cut.rsplit(" ", 1)[0]
            lines.append(cut)
            raw = raw[len(cut):].lstrip()
        lines.append(raw)

    for line in lines:
        if y < 60:
            c.showPage()
            y = height - 60
            c.setFont("Helvetica", 11)
        c.drawString(x, y, line)
        y -= 14

    c.save()
    return bio.getvalue()


# ==========================================================
# UI + lógica
# ==========================================================
audio_file = st.file_uploader(
    "Sube audio (mp3, m4a, wav, webm, mp4)",
    type=["mp3", "m4a", "wav", "webm", "mp4", "mpeg", "mpga"],
)

if audio_file is not None:
    st.audio(audio_file)

btn = st.button("Transcribir", type="primary", disabled=(audio_file is None), use_container_width=True)

if btn and audio_file is not None:
    t0 = time.time()
    audio_bytes = audio_file.read()

    if len(audio_bytes) > MAX_FILE_MB * 1024 * 1024:
        st.warning(f"El archivo supera {MAX_FILE_MB} MB. Por favor divide el audio o usa un formato más comprimido.")
        st.stop()

    with st.spinner("Transcribiendo…"):
        transcript_text, _payload = transcribe_openai(audio_bytes, audio_file.name)

    st.success("Listo ✅")
    st.caption(f"Tiempo: {round(time.time() - t0, 2)} s")

    st.subheader("Transcripción")
    st.text_area("Resultado", value=transcript_text, height=360)

    st.subheader("Exportar")
    c1, c2, c3 = st.columns(3)

    with c1:
        st.download_button(
            "TXT",
            data=transcript_text.encode("utf-8"),
            file_name="transcripcion.txt",
            mime="text/plain",
            use_container_width=True,
        )

    with c2:
        docx_bytes = to_docx_bytes("Transcripción", transcript_text)
        if docx_bytes is None:
            st.info("Para exportar DOCX agrega `python-docx` a requirements.txt.")
        else:
            st.download_button(
                "DOCX",
                data=docx_bytes,
                file_name="transcripcion.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

    with c3:
        pdf_bytes = to_pdf_bytes("Transcripción", transcript_text)
        if pdf_bytes is None:
            st.info("Para exportar PDF agrega `reportlab` a requirements.txt.")
        else:
            st.download_button(
                "PDF",
                data=pdf_bytes,
                file_name="transcripcion.pdf",
                mime="application/pdf",
                use_container_width=True,
            )

