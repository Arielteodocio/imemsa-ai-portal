import base64
import io
import json
import os
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple
from utils_portal_auth import require_login_redirect

import pandas as pd
import requests
import streamlit as st
from imemsa_ui import render_title
from PIL import Image

# ==========================================================
# PÁGINA: Documentos (OCR + extracción)
# - Compatible con el proyecto base (login en app.py)
# - SIN st.set_page_config() (solo en app.py)
# - SIN st.switch_page() / st.rerun() (evita loops)
# ==========================================================

# --------- Login guard (misma llave que app.py del proyecto base)
#def require_login() -> None:
#    if not st.session_state.get("auth", False):
#        st.error("🔒 Inicia sesión para usar esta herramienta.")
#        if hasattr(st, "page_link"):
#            st.page_link("app.py", label="Ir al Login", icon="🔐", use_container_width=True)
#        st.stop()

require_login_redirect()

#require_login()

# --------- UI Header
render_title('📄 Documentos', 'Sube un PDF o imagen y obtén texto (OCR) y extracción estructurada.')

if hasattr(st, "page_link"):
    st.page_link("app.py", label="⬅️ Volver al Portafolio", icon="🏠", use_container_width=True)


with st.expander("🔒 Privacidad (cómo funciona)", expanded=False):
    st.write(
        "- El documento se envía a la IA para **leerlo** y devolver texto/campos.\n"
        "- La app no guarda el archivo.\n"
        "- Puedes exportar el resultado (TXT/JSON/CSV y opcional DOCX/PDF)."
    )

# ==========================================================
# Config y helpers
# ==========================================================
MODEL = "gpt-4o-mini"  # multimodal
MAX_FILE_MB = 25


def _get_openai_api_key() -> Optional[str]:
    return (
        st.secrets.get("OPENAI_API_KEY") if hasattr(st, "secrets") else None
    ) or os.getenv("OPENAI_API_KEY")


def _b64_data_url(img_bytes: bytes, mime: str) -> str:
    b64 = base64.b64encode(img_bytes).decode("utf-8")
    return f"data:{mime};base64,{b64}"


def _img_to_png_bytes(file) -> Tuple[bytes, str]:
    """Normaliza imágenes a PNG (mejor lectura)."""
    img = Image.open(file).convert("RGB")

    # Downscale suave para evitar requests gigantes (mantiene proporción)
    max_side = 1800
    w, h = img.size
    scale = min(1.0, max_side / max(w, h))
    if scale < 1.0:
        img = img.resize((int(w * scale), int(h * scale)))

    bio = io.BytesIO()
    img.save(bio, format="PNG")
    return bio.getvalue(), "image/png"


def _pdf_to_png_pages(pdf_bytes: bytes, dpi: int = 150, max_pages: int = 6) -> List[Tuple[bytes, str, str]]:
    """
    Convierte PDF → lista de páginas PNG.
    Requiere 'pymupdf' (fitz). Si no está instalado, muestra instrucción.
    """
    try:
        import fitz  # PyMuPDF
    except Exception:
        st.error("Para leer PDFs necesito `pymupdf` (PyMuPDF).")
        st.info("Agrega `pymupdf` a tu requirements.txt y redeploy.")
        st.stop()

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages = []
    n = min(len(doc), max_pages)
    for i in range(n):
        page = doc[i]
        pix = page.get_pixmap(dpi=dpi)
        img_bytes = pix.tobytes("png")
        pages.append((img_bytes, "image/png", f"page_{i+1}.png"))
    if len(doc) > max_pages:
        st.warning(f"El PDF tiene {len(doc)} páginas. Procesé solo las primeras {max_pages}.")
    return pages


def _openai_chat(messages: List[Dict[str, Any]], temperature: float = 0.2, timeout: int = 180) -> str:
    api_key = _get_openai_api_key()
    if not api_key:
        st.error("Falta la API Key.")
        st.info(
            "En Streamlit Cloud agrega un Secret llamado **OPENAI_API_KEY** "
            "o define la variable de entorno `OPENAI_API_KEY`."
        )
        st.stop()

    url = "https://api.openai.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {"model": MODEL, "messages": messages, "temperature": temperature}

    resp = requests.post(url, headers=headers, json=payload, timeout=timeout)
    if resp.status_code >= 400:
        st.error(f"Error del servicio (HTTP {resp.status_code}).")
        try:
            st.code(resp.json())
        except Exception:
            st.code(resp.text[:2000])
        st.stop()

    data = resp.json()
    return data["choices"][0]["message"]["content"].strip()


def _extract_json(text: str) -> Dict[str, Any]:
    text = (text or "").strip()
    try:
        return json.loads(text)
    except Exception:
        # intenta extraer objeto JSON del texto
        import re

        m = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not m:
            raise ValueError("No se encontró JSON válido en la respuesta.")
        return json.loads(m.group(0))


@dataclass
class DocResult:
    full_text: str
    fields: Dict[str, Any]


def ocr_and_extract(images: List[Tuple[bytes, str, str]], doc_type: str) -> DocResult:
    """
    En un solo paso:
    - OCR (texto completo)
    - Extracción estructurada (campos + items)
    """
    schema = {
        "full_text": "string",
        "fields": {
            "tipo_documento": "Factura | Orden de compra | Reporte | Checklist | Otro | ''",
            "folio": "string o ''",
            "fecha": "string o ''",
            "empresa": "string o ''",
            "proveedor": "string o ''",
            "total": "string o ''",
            "moneda": "string o ''",
            "items": [
                {
                    "descripcion": "string",
                    "cantidad": "string o ''",
                    "unidad": "string o ''",
                    "precio_unitario": "string o ''",
                    "importe": "string o ''",
                }
            ],
            "notas": "string o ''",
        },
    }

    instruction = (
        "Eres un asistente experto en OCR y extracción de datos.\n"
        "Devuelve SOLO un JSON válido (sin markdown, sin texto extra).\n"
        "Reglas:\n"
        "1) No inventes datos: si no aparece, deja '' o lista vacía.\n"
        "2) Conserva números, unidades y el texto tal cual.\n"
        "3) Si el documento es tabla, intenta extraer items.\n"
        f"4) Si el usuario indicó tipo_documento='{doc_type}', úsalo como pista.\n"
        "Esquema EXACTO:\n"
        f"{json.dumps(schema, ensure_ascii=False)}"
    )

    content: List[Dict[str, Any]] = [{"type": "text", "text": instruction}]
    for img_bytes, mime, _name in images:
        content.append({"type": "image_url", "image_url": {"url": _b64_data_url(img_bytes, mime)}})

    messages = [{"role": "user", "content": content}]
    raw = _openai_chat(messages, temperature=0.1, timeout=240)
    obj = _extract_json(raw)

    full_text = (obj.get("full_text") or "").strip()
    fields = obj.get("fields") or {}
    if not isinstance(fields, dict):
        fields = {}

    # normaliza items si existen
    items = fields.get("items", [])
    if not isinstance(items, list):
        fields["items"] = []
    else:
        # limpia registros vacíos
        clean_items = []
        for it in items:
            if not isinstance(it, dict):
                continue
            desc = (it.get("descripcion") or "").strip()
            if not desc:
                continue
            clean_items.append(
                {
                    "descripcion": desc,
                    "cantidad": (it.get("cantidad") or "").strip(),
                    "unidad": (it.get("unidad") or "").strip(),
                    "precio_unitario": (it.get("precio_unitario") or "").strip(),
                    "importe": (it.get("importe") or "").strip(),
                }
            )
        fields["items"] = clean_items

    return DocResult(full_text=full_text, fields=fields)


# ==========================================================
# Export helpers (opcionales)
# ==========================================================
def to_docx_bytes(title: str, body: str) -> Optional[bytes]:
    try:
        from docx import Document
    except Exception:
        return None

    doc = Document()
    doc.add_heading(title, level=1)
    for para in body.split("\n"):
        doc.add_paragraph(para)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def to_pdf_bytes(title: str, body: str) -> Optional[bytes]:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas
    except Exception:
        return None

    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=LETTER)
    _, height = LETTER

    x = 50
    y = height - 60
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, title)
    y -= 25
    c.setFont("Helvetica", 11)

    max_chars = 95
    lines: List[str] = []
    for raw in body.split("\n"):
        raw = raw.rstrip()
        if not raw:
            lines.append("")
            continue
        while len(raw) > max_chars:
            cut = raw[:max_chars]
            if " " in cut:
                cut = cut.rsplit(" ", 1)[0]
            lines.append(cut)
            raw = raw[len(cut):].lstrip()
        lines.append(raw)

    for line in lines:
        if y < 60:
            c.showPage()
            y = height - 60
            c.setFont("Helvetica", 11)
        c.drawString(x, y, line)
        y -= 14

    c.save()
    return bio.getvalue()


# ==========================================================
# UI
# ==========================================================
with st.sidebar:
    st.markdown("### Opciones")
    max_pages = st.slider("Páginas máximas (PDF)", min_value=1, max_value=12, value=6, step=1)
    dpi = st.slider("Calidad PDF → imagen (DPI)", min_value=100, max_value=220, value=150, step=10)

doc_type = st.selectbox(
    "Tipo de documento (opcional)",
    ["Auto", "Factura", "Orden de compra", "Reporte", "Checklist", "Otro"],
    index=0,
)

uploaded = st.file_uploader("Sube PDF o imagen", type=["pdf", "png", "jpg", "jpeg"])
st.caption("Tip: si el PDF es muy pesado, prueba bajar dpi o limitar páginas.")

btn = st.button("Procesar documento", type="primary", disabled=(uploaded is None), use_container_width=True)

if uploaded is not None:
    size_mb = len(uploaded.getvalue()) / (1024 * 1024)
    if size_mb > MAX_FILE_MB:
        st.warning(f"El archivo pesa {size_mb:.1f} MB. Recomendado: ≤ {MAX_FILE_MB} MB.")

if btn and uploaded is not None:
    file_bytes = uploaded.getvalue()
    ext = uploaded.name.lower().split(".")[-1]

    hint = "" if doc_type == "Auto" else doc_type

    try:
        with st.spinner("Procesando…"):
            images: List[Tuple[bytes, str, str]] = []

            if ext == "pdf":
                images = _pdf_to_png_pages(file_bytes, dpi=dpi, max_pages=max_pages)
            else:
                img_bytes, mime = _img_to_png_bytes(uploaded)
                images = [(img_bytes, mime, uploaded.name)]

            # Preview
            with st.expander("Vista previa", expanded=False):
                for b, _m, name in images[:4]:
                    st.image(b, caption=name, use_column_width=True)

            result = ocr_and_extract(images, doc_type=hint)

        st.success("Listo ✅")

        st.subheader("Texto (OCR)")
        st.text_area("Texto extraído", value=result.full_text, height=320)

        st.subheader("Extracción estructurada")
        st.json(result.fields)

        items = result.fields.get("items", [])
        df_items = pd.DataFrame(items) if isinstance(items, list) else pd.DataFrame()
        if not df_items.empty:
            st.subheader("Items detectados")
            st.dataframe(df_items, use_container_width=True)

        st.divider()
        st.subheader("Exportar")

        txt_out = (result.full_text or "").strip()
        json_out = json.dumps(result.fields, ensure_ascii=False, indent=2)

        c1, c2, c3, c4, c5 = st.columns(5)

        with c1:
            st.download_button("TXT", txt_out.encode("utf-8"), "documento_ocr.txt", "text/plain", use_container_width=True)

        with c2:
            docx_bytes = to_docx_bytes("Documento (OCR)", txt_out)
            if docx_bytes is None:
                st.info("DOCX: agrega `python-docx` a requirements.txt.")
            else:
                st.download_button(
                    "DOCX",
                    docx_bytes,
                    "documento_ocr.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

        with c3:
            pdf_bytes = to_pdf_bytes("Documento (OCR)", txt_out)
            if pdf_bytes is None:
                st.info("PDF: agrega `reportlab` a requirements.txt.")
            else:
                st.download_button("PDF", pdf_bytes, "documento_ocr.pdf", "application/pdf", use_container_width=True)

        with c4:
            st.download_button("JSON", json_out.encode("utf-8"), "campos.json", "application/json", use_container_width=True)

        with c5:
            if not df_items.empty:
                st.download_button(
                    "CSV (items)",
                    df_items.to_csv(index=False).encode("utf-8"),
                    "items.csv",
                    "text/csv",
                    use_container_width=True,
                )
            else:
                kv = pd.DataFrame([{"campo": k, "valor": v} for k, v in result.fields.items()])
                st.download_button(
                    "CSV (campos)",
                    kv.to_csv(index=False).encode("utf-8"),
                    "campos.csv",
                    "text/csv",
                    use_container_width=True,
                )

    except Exception as e:
        st.error("Ocurrió un error al procesar el documento.")
        with st.expander("🛠️ Detalle técnico (solo admin)", expanded=False):
            st.exception(e)
