import io
import os
from typing import Optional, Tuple

import requests
import streamlit as st
from imemsa_ui import render_title
from utils_portal_auth import require_login_redirect

# ==========================================================
# PÁGINA: Traducción (Texto → Texto)
# - Compatible con el proyecto base (login en app.py)
# - SIN st.set_page_config() (solo en app.py)
# - SIN st.switch_page() / st.rerun() (evita loops)
# ==========================================================

# --------- Login guard (misma llave que app.py del proyecto base)
#def require_login() -> None:
#    if not st.session_state.get("auth", False):
#        st.error("🔒 Inicia sesión para usar esta herramienta.")
#        if hasattr(st, "page_link"):
#            st.page_link("app.py", label="Ir al Login", icon="🔐", use_container_width=True)
#        st.stop()
require_login_redirect()


#require_login()

# --------- UI Header
render_title('🌐 Traducción', 'Traduce entre Español ↔ Inglés manteniendo formato, números y términos técnicos.')

if hasattr(st, "page_link"):
    st.page_link("app.py", label="⬅️ Volver al Portafolio", icon="🏠", use_container_width=True)


with st.expander("Consejos de traducción (industrial/técnico)", expanded=False):
    st.write(
        "- Mantén siglas (IMEMSA, ERP, SLA), unidades (mm, kg) y números.\n"
        "- Conserva listas, viñetas y saltos de línea.\n"
        "- Si hay términos que no deben traducirse (p. ej. *gel coat*, *T-top*), agrégalos al glosario."
    )


# ==========================================================
# OpenAI Chat Completions via HTTP (sin SDK)
# ==========================================================
MODEL = "gpt-4o-mini"

def _get_openai_api_key() -> Optional[str]:
    return (
        st.secrets.get("OPENAI_API_KEY") if hasattr(st, "secrets") else None
    ) or os.getenv("OPENAI_API_KEY")


def translate_text(text: str, direction: str, tone: str, glossary: str) -> str:
    api_key = _get_openai_api_key()
    if not api_key:
        st.error("Falta la API Key.")
        st.info(
            "En Streamlit Cloud agrega un Secret llamado **OPENAI_API_KEY** "
            "o define la variable de entorno `OPENAI_API_KEY`."
        )
        st.stop()

    # Dirección
    if direction == "ES → EN":
        src, tgt = "español", "inglés"
    else:
        src, tgt = "inglés", "español"

    # Glosario (opcional)
    glossary = (glossary or "").strip()
    glossary_block = ""
    if glossary:
        glossary_block = (
            "\n\nGlosario (respétalo estrictamente):\n"
            "- Si aparece exactamente uno de estos términos, mantén o traduce según se indique.\n"
            f"{glossary}\n"
        )

    system = (
        "Eres un traductor profesional. Devuelve SOLO la traducción final, sin explicaciones. "
        "Conserva el formato (saltos de línea, viñetas, tablas simples), números, unidades y nombres propios. "
        "No inventes información."
    )

    user = (
        f"Traduce del {src} al {tgt} con tono {tone}.\n"
        "Reglas:\n"
        "1) Mantén siglas y términos técnicos.\n"
        "2) Conserva el formato y los saltos de línea.\n"
        "3) Si hay texto dentro de comillas, tradúcelo respetando las comillas.\n"
        "4) Devuelve solo la traducción.\n"
        f"{glossary_block}\n"
        "TEXTO:\n"
        f"{text}"
    )

    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": MODEL,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.2,
    }

    resp = requests.post(url, headers=headers, json=payload, timeout=120)

    if resp.status_code >= 400:
        st.error(f"Error en traducción (HTTP {resp.status_code}).")
        try:
            st.code(resp.json())
        except Exception:
            st.code(resp.text[:2000])
        st.stop()

    data = resp.json()
    try:
        return data["choices"][0]["message"]["content"].strip()
    except Exception:
        st.error("Respuesta inesperada del servicio de traducción.")
        st.code(data)
        st.stop()


# ==========================================================
# Export helpers (opcionales)
# ==========================================================
def to_docx_bytes(title: str, body: str) -> Optional[bytes]:
    try:
        from docx import Document
    except Exception:
        return None

    doc = Document()
    doc.add_heading(title, level=1)
    for para in body.split("\n"):
        doc.add_paragraph(para)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def to_pdf_bytes(title: str, body: str) -> Optional[bytes]:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas
    except Exception:
        return None

    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=LETTER)
    width, height = LETTER

    x = 50
    y = height - 60
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, title)
    y -= 25
    c.setFont("Helvetica", 11)

    max_chars = 95
    lines = []
    for raw in body.split("\n"):
        raw = raw.rstrip()
        if not raw:
            lines.append("")
            continue
        while len(raw) > max_chars:
            cut = raw[:max_chars]
            if " " in cut:
                cut = cut.rsplit(" ", 1)[0]
            lines.append(cut)
            raw = raw[len(cut):].lstrip()
        lines.append(raw)

    for line in lines:
        if y < 60:
            c.showPage()
            y = height - 60
            c.setFont("Helvetica", 11)
        c.drawString(x, y, line)
        y -= 14

    c.save()
    return bio.getvalue()


# ==========================================================
# UI
# ==========================================================
col1, col2 = st.columns([1, 1])
with col1:
    direction = st.radio("Dirección", ["ES → EN", "EN → ES"], horizontal=True)
with col2:
    tone = st.selectbox("Tono", ["técnico", "formal", "neutral"], index=0)

glossary = st.text_area(
    "Glosario (opcional)",
    placeholder="Ejemplos:\n- gel coat -> gel coat (no traducir)\n- borda -> gunwale\n- patín -> strake\n",
    height=110,
)

text = st.text_area(
    "Texto a traducir",
    height=260,
    placeholder="Pega aquí el texto…",
)

btn = st.button("Traducir", type="primary", disabled=(not text.strip()), use_container_width=True)

if btn:
    # guard simple contra entradas enormes (para evitar timeouts)
    if len(text) > 18000:
        st.warning("El texto es muy largo. Divide en partes (recomendado: < 18,000 caracteres).")
        st.stop()

    with st.spinner("Traduciendo…"):
        result = translate_text(text=text, direction=direction, tone=tone, glossary=glossary)

    st.success("Listo ✅")
    st.subheader("Resultado")
    st.text_area("Traducción", value=result, height=320)

    st.subheader("Exportar")
    c1, c2, c3 = st.columns(3)

    with c1:
        st.download_button(
            "TXT",
            data=result.encode("utf-8"),
            file_name="traduccion.txt",
            mime="text/plain",
            use_container_width=True,
        )

    with c2:
        docx_bytes = to_docx_bytes("Traducción", result)
        if docx_bytes is None:
            st.info("Para exportar DOCX agrega `python-docx` a requirements.txt.")
        else:
            st.download_button(
                "DOCX",
                data=docx_bytes,
                file_name="traduccion.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

    with c3:
        pdf_bytes = to_pdf_bytes("Traducción", result)
        if pdf_bytes is None:
            st.info("Para exportar PDF agrega `reportlab` a requirements.txt.")
        else:
            st.download_button(
                "PDF",
                data=pdf_bytes,
                file_name="traduccion.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
