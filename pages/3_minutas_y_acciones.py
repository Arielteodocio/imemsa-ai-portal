import io
import json
import os
import re
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple
from utils_ui import render_title

import pandas as pd
import requests
import streamlit as st

# ==========================================================
# PÁGINA: Minutas y acciones (Transcripción → Minuta + Acciones)
# - Compatible con el proyecto base (login en app.py)
# - SIN st.set_page_config() (solo en app.py)
# - SIN st.switch_page() / st.rerun() (evita loops)
# ==========================================================

# --------- Login guard (misma llave que app.py del proyecto base)
def require_login() -> None:
    if not st.session_state.get("auth", False):
        st.error("🔒 Inicia sesión para usar esta herramienta.")
        if hasattr(st, "page_link"):
            st.page_link("app.py", label="Ir al Login", icon="🔐", use_container_width=True)
        st.stop()


require_login()

# --------- UI Header
st.markdown("# 📝 Minutas y acciones")
st.caption(
    "Pega la transcripción de una reunión y genera una minuta estructurada con acuerdos y acciones "
    "(responsables y fechas si aparecen en el texto)."
)

if hasattr(st, "page_link"):
    st.page_link("app.py", label="⬅️ Volver al Portafolio", icon="🏠", use_container_width=True)

st.divider()

with st.expander("🔒 Privacidad (cómo funciona)", expanded=False):
    st.write(
        "- El texto se envía al servicio de IA para estructurarlo.\n"
        "- Esta app no guarda tu transcripción ni la minuta.\n"
        "- Puedes exportar los resultados en TXT/CSV/Excel (y DOCX/PDF si tienes dependencias)."
    )

# ==========================================================
# OpenAI Chat Completions via HTTP (sin SDK)
# ==========================================================
MODEL = "gpt-4o-mini"

def _get_openai_api_key() -> Optional[str]:
    return (
        st.secrets.get("OPENAI_API_KEY") if hasattr(st, "secrets") else None
    ) or os.getenv("OPENAI_API_KEY")


def _extract_json(text: str) -> Dict[str, Any]:
    """
    El modelo debe devolver JSON, pero si trae texto extra,
    intentamos extraer el objeto JSON del contenido.
    """
    text = (text or "").strip()

    # Caso ideal: es JSON puro
    try:
        return json.loads(text)
    except Exception:
        pass

    # Busca el primer { ... último }
    m = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if not m:
        raise ValueError("No se encontró un objeto JSON en la respuesta.")
    return json.loads(m.group(0))


@dataclass
class MinutesResult:
    title: str
    summary: str
    agreements: List[str]
    actions: List[Dict[str, Any]]


def generate_minutes(transcript: str, tone: str = "técnico") -> MinutesResult:
    api_key = _get_openai_api_key()
    if not api_key:
        st.error("Falta la API Key.")
        st.info(
            "En Streamlit Cloud agrega un Secret llamado **OPENAI_API_KEY** "
            "o define la variable de entorno `OPENAI_API_KEY`."
        )
        st.stop()

    system = (
        "Eres un asistente experto en redacción de minutas para contexto industrial y administrativo. "
        "Devuelve SOLO un JSON válido (sin markdown, sin texto extra). "
        "Si algún dato no existe en la transcripción, usa cadena vacía o null; no inventes."
    )

    schema = {
        "title": "string (título corto de la reunión)",
        "summary": "string (resumen ejecutivo 5-10 líneas)",
        "agreements": ["string (acuerdo)"],
        "actions": [
            {
                "accion": "string",
                "responsable": "string o ''",
                "fecha_compromiso": "string o '' (si se menciona, formato libre)",
                "prioridad": "alta|media|baja o ''",
                "area": "string o ''",
                "notas": "string o ''",
            }
        ],
    }

    user = (
        f"Genera una minuta en español con tono {tone} a partir de la siguiente transcripción.\n"
        "Reglas:\n"
        "1) No inventes responsables, fechas o prioridades.\n"
        "2) Si detectas varias acciones, separa una por registro.\n"
        "3) Mantén términos técnicos (IMEMSA, fibra de vidrio, gel coat, infusión al vacío, T-top).\n"
        "4) Devuelve JSON con este esquema EXACTO:\n"
        f"{json.dumps(schema, ensure_ascii=False)}\n\n"
        "TRANSCRIPCIÓN:\n"
        f"{transcript}"
    )

    url = "https://api.openai.com/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {
        "model": MODEL,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.2,
    }

    resp = requests.post(url, headers=headers, json=payload, timeout=180)
    if resp.status_code >= 400:
        st.error(f"Error al generar minuta (HTTP {resp.status_code}).")
        try:
            st.code(resp.json())
        except Exception:
            st.code(resp.text[:2000])
        st.stop()

    data = resp.json()
    content = data["choices"][0]["message"]["content"]
    obj = _extract_json(content)

    title = (obj.get("title") or "").strip() or "Minuta"
    summary = (obj.get("summary") or "").strip()
    agreements = obj.get("agreements") or []
    actions = obj.get("actions") or []

    # Normaliza tipos
    if not isinstance(agreements, list):
        agreements = []
    agreements = [str(x).strip() for x in agreements if str(x).strip()]

    if not isinstance(actions, list):
        actions = []

    # Asegura columnas esperadas
    cleaned_actions = []
    for a in actions:
        a = a if isinstance(a, dict) else {}
        cleaned_actions.append(
            {
                "accion": (a.get("accion") or "").strip(),
                "responsable": (a.get("responsable") or "").strip(),
                "fecha_compromiso": (a.get("fecha_compromiso") or "").strip(),
                "prioridad": (a.get("prioridad") or "").strip(),
                "area": (a.get("area") or "").strip(),
                "notas": (a.get("notas") or "").strip(),
            }
        )

    # Filtra acciones vacías
    cleaned_actions = [a for a in cleaned_actions if a.get("accion")]

    return MinutesResult(
        title=title,
        summary=summary,
        agreements=agreements,
        actions=cleaned_actions,
    )


# ==========================================================
# Export helpers (opcionales)
# ==========================================================
def to_docx_bytes(title: str, body: str) -> Optional[bytes]:
    try:
        from docx import Document
    except Exception:
        return None

    doc = Document()
    doc.add_heading(title, level=1)
    for para in body.split("\n"):
        doc.add_paragraph(para)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def to_pdf_bytes(title: str, body: str) -> Optional[bytes]:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas
    except Exception:
        return None

    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=LETTER)
    width, height = LETTER

    x = 50
    y = height - 60
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, title)
    y -= 25
    c.setFont("Helvetica", 11)

    max_chars = 95
    lines = []
    for raw in body.split("\n"):
        raw = raw.rstrip()
        if not raw:
            lines.append("")
            continue
        while len(raw) > max_chars:
            cut = raw[:max_chars]
            if " " in cut:
                cut = cut.rsplit(" ", 1)[0]
            lines.append(cut)
            raw = raw[len(cut):].lstrip()
        lines.append(raw)

    for line in lines:
        if y < 60:
            c.showPage()
            y = height - 60
            c.setFont("Helvetica", 11)
        c.drawString(x, y, line)
        y -= 14

    c.save()
    return bio.getvalue()


def actions_to_xlsx_bytes(df: pd.DataFrame) -> Optional[bytes]:
    try:
        from openpyxl import Workbook
        from openpyxl.utils.dataframe import dataframe_to_rows
    except Exception:
        return None

    wb = Workbook()
    ws = wb.active
    ws.title = "Acciones"

    for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), start=1):
        ws.append(row)

    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()


# ==========================================================
# UI
# ==========================================================
tone = st.selectbox("Tono de redacción", ["técnico", "formal", "neutral"], index=0)

transcript = st.text_area(
    "Transcripción",
    height=320,
    placeholder="Pega aquí la transcripción completa de la reunión…",
)

btn = st.button("Generar minuta", type="primary", disabled=(not transcript.strip()), use_container_width=True)

if btn:
    # guard simple contra entradas enormes (evita timeouts)
    if len(transcript) > 35000:
        st.warning("La transcripción es muy larga. Divide en partes (recomendado: < 35,000 caracteres).")
        st.stop()

    with st.spinner("Generando minuta…"):
        result = generate_minutes(transcript.strip(), tone=tone)

    st.success("Listo ✅")

    # ---- Presentación
    st.subheader(result.title)

    st.markdown("### Resumen")
    st.write(result.summary if result.summary else "_(Sin resumen)_")

    st.markdown("### Acuerdos")
    if result.agreements:
        for i, a in enumerate(result.agreements, start=1):
            st.write(f"{i}. {a}")
    else:
        st.info("No se detectaron acuerdos explícitos.")

    st.markdown("### Acciones")
    expected_cols = ["accion", "responsable", "fecha_compromiso", "prioridad", "area", "notas"]
    df = pd.DataFrame(result.actions, columns=expected_cols) if result.actions else pd.DataFrame(columns=expected_cols)

    edited_df = st.data_editor(
        df,
        use_container_width=True,
        num_rows="dynamic",
        key="actions_editor",
    )

    # ---- Exportables
    st.divider()
    st.subheader("Exportar")

    # Texto consolidado
    lines: List[str] = []
    lines.append(result.title)
    lines.append("")
    lines.append("RESUMEN")
    lines.append(result.summary or "")
    lines.append("")
    lines.append("ACUERDOS")
    if result.agreements:
        for i, a in enumerate(result.agreements, start=1):
            lines.append(f"{i}. {a}")
    else:
        lines.append("No se detectaron acuerdos.")
    lines.append("")
    lines.append("ACCIONES")
    if not edited_df.empty:
        for _, row in edited_df.iterrows():
            lines.append(
                f"- Acción: {row.get('accion','')}\n"
                f"  Responsable: {row.get('responsable','')}\n"
                f"  Fecha compromiso: {row.get('fecha_compromiso','')}\n"
                f"  Prioridad: {row.get('prioridad','')}\n"
                f"  Área: {row.get('area','')}\n"
                f"  Notas: {row.get('notas','')}"
            )
    else:
        lines.append("No se detectaron acciones.")
    txt_out = "\n".join(lines).strip()

    c1, c2, c3, c4, c5 = st.columns(5)

    with c1:
        st.download_button(
            "TXT",
            data=txt_out.encode("utf-8"),
            file_name="minuta.txt",
            mime="text/plain",
            use_container_width=True,
        )

    with c2:
        docx_bytes = to_docx_bytes(result.title, txt_out)
        if docx_bytes is None:
            st.info("DOCX: agrega `python-docx` a requirements.txt.")
        else:
            st.download_button(
                "DOCX",
                data=docx_bytes,
                file_name="minuta.docx",
                use_container_width=True,
            )

    with c3:
        pdf_bytes = to_pdf_bytes(result.title, txt_out)
        if pdf_bytes is None:
            st.info("PDF: agrega `reportlab` a requirements.txt.")
        else:
            st.download_button(
                "PDF",
                data=pdf_bytes,
                file_name="minuta.pdf",
                mime="application/pdf",
                use_container_width=True,
            )

    with c4:
        st.download_button(
            "CSV (acciones)",
            data=edited_df.to_csv(index=False).encode("utf-8"),
            file_name="acciones.csv",
            mime="text/csv",
            use_container_width=True,
        )

    with c5:
        xlsx_bytes = actions_to_xlsx_bytes(edited_df)
        if xlsx_bytes is None:
            st.info("Excel: agrega `openpyxl` a requirements.txt.")
        else:
            st.download_button(
                "Excel (acciones)",
                data=xlsx_bytes,
                file_name="acciones.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )

