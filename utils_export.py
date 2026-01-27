from io import BytesIO
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from docx import Document


def to_docx_bytes(title: str, text: str) -> bytes:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for para in (text or "").split("\n"):
        doc.add_paragraph(para)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def to_pdf_bytes(title: str, text: str) -> bytes:
    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=LETTER)
    width, height = LETTER

    y = height - 72
    if title:
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, y, title[:120])
        y -= 24

    c.setFont("Helvetica", 11)
    for line in (text or "").split("\n"):
        while len(line) > 95:
            c.drawString(72, y, line[:95])
            line = line[95:]
            y -= 14
            if y < 72:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - 72
        c.drawString(72, y, line)
        y -= 14
        if y < 72:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - 72

    c.save()
    return bio.getvalue()
